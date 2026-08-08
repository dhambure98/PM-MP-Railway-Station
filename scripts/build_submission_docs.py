"""
Generate the two EEI6373 mini-project submission Word documents:
  1. docs/Submission_1_Problem_and_Dataset.docx  (high-level problem + dataset)
  2. docs/Submission_2_Final_Report.docx          (full final report + figures)
"""

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
DATA = ROOT / "data" / "passenger_flow_peak_hours.csv"
OUTPUTS = ROOT / "outputs"
DOCS = ROOT / "docs"
GIT_URL = "https://github.com/dhambure98/PM-MP-Railway-Station.git"

STUDENT = "Akila Dhambure Liyanage"
REG_NO = "320267120"
COURSE = "EEI6373 – Performance Modelling"
TOPIC = "Performance Modelling of Passenger Flow at a Major Railway Station during Peak Hours"
TITLE = "Railway Station Passenger Flow"

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ModuleNotFoundError:
    sys.exit("python-docx is required: pip install python-docx")


def set_cell(cell, text, bold=False, center=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def make_table(doc, headers, rows, widths=None, col_center=True):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        set_cell(table.rows[0].cells[j], h, bold=True, center=True)
    for row in rows:
        cells = table.add_row().cells
        for j, val in enumerate(row):
            set_cell(cells[j], str(val), center=col_center)
    if widths:
        for j, w in enumerate(widths):
            for row in table.rows:
                row.cells[j].width = Inches(w)
    doc.add_paragraph()
    return table


def add_title_block(doc, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(20)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.italic = True
    run.font.size = Pt(13)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(COURSE).font.size = Pt(11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"{STUDENT}  |  {REG_NO}").font.size = Pt(11)
    doc.add_paragraph()


def build_submission_1():
    doc = Document()
    add_title_block(doc, "Submission 1 – High-Level Problem and Data Set")

    doc.add_heading("1. Identification of the Complex System", level=1)
    doc.add_paragraph(
        "This mini project studies the passenger flow through the ticket concourse of a "
        "major railway station during the morning peak period (07:00 – 09:00). The station "
        "is modelled as a complex service system in which passengers must pass through a "
        "chain of stages: ticket acquisition, entry gates, platform access and boarding. "
        "The system is complex because the stages are coupled, arrivals are stochastic, "
        "demand is sharply concentrated in a short window, and the performance of one "
        "stage directly constrains the stages that follow."
    )
    doc.add_paragraph(
        "Two passenger classes are distinguished. Normal ticket passengers must queue at "
        "one of three ticket counters before proceeding to the gates, whereas Online "
        "ticket passengers bypass the counters entirely and go straight to the gates. "
        "This mixed demand pattern creates an uneven, time-varying load on the counter "
        "subsystem and makes the ticket counters the natural bottleneck candidate."
    )

    doc.add_heading("1.1 Why the System Is Complex", level=2)
    for item in [
        "Stochastic behaviour – passenger arrivals are random, and service times vary passenger by passenger.",
        "Coupled stages – congestion at the counters propagates into gate and boarding queues.",
        "Class-dependent demand – Normal and Online passengers place very different loads on the counters.",
        "Peak concentration – almost all demand occurs within two hours, producing a heavy transient.",
        "Limited resources – fixed counter capacity and gate capacity must be dimensioned to absorb the peak.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2. Performance Objectives", level=1)
    doc.add_paragraph(
        "The project defines the following quantifiable performance objectives for the "
        "system, consistent with queueing-theoretic measures:"
    )
    doc.add_paragraph(
        "O1 – Identify the bottleneck stage of the passenger flow and quantify its "
        "utilisation under normal and busy demand.",
        style="List Number",
    )
    doc.add_paragraph(
        "O2 – Estimate the expected waiting time at the ticket counters (Wq), the "
        "probability that a passenger has to wait (P(wait)), and the expected queue "
        "length (Lq) under the baseline peak demand.",
        style="List Number",
    )
    doc.add_paragraph(
        "O3 – Determine whether the current configuration (c = 3 counters, μ = 40 "
        "passengers/hour) is stable, and find the arrival-rate capacity limit.",
        style="List Number",
    )
    doc.add_paragraph(
        "O4 – Evaluate what-if interventions (additional counters, faster service) and "
        "show how they reduce waiting under high load.",
        style="List Number",
    )
    doc.add_paragraph(
        "O5 – Validate the analytical model (M/M/c, Erlang-C, Little's Law) against a "
        "passenger-level simulation to confirm the estimates are trustworthy.",
        style="List Number",
    )

    doc.add_heading("3. Data Set", level=1)
    doc.add_paragraph(
        "The data set is a passenger-level record of the 07:00 – 09:00 peak period "
        "generated by a FIFO M/M/c discrete-event simulation (c = 3, μ = 40/hour). "
        "It contains 400 passengers: 180 Normal (45.0%) and 220 Online (55.0%), "
        "corresponding to arrival rates of λ = 90/hour and λ = 110/hour respectively. "
        "The file is stored in the repository at data/passenger_flow_peak_hours.csv."
    )
    doc.add_paragraph(
        "Important: the passenger-level data are simulated, not real station "
        "observations. They are produced so that the data set validates the analytical "
        "model used in the project."
    )

    doc.add_heading("3.1 Columns", level=2)
    make_table(
        doc,
        ["Column", "Description"],
        [
            ["Passenger_ID", "Unique passenger identifier (P001 … P400)"],
            ["Arrival_Time", "Time of arrival into the concourse (07:00:00 – 09:00:00)"],
            ["Ticket_Type", "Passenger class: Normal (uses counters) or Online (bypasses counters)"],
            ["Counter_Wait_Min", "Waiting time at the ticket counters (min); 0 for Online passengers"],
            ["Counter_Service_Min", "Service time at the ticket counters (min); 0 for Online passengers"],
            ["Gate_Wait_Min", "Waiting time at the entry gates (min)"],
            ["Gate_Service_Min", "Service time at the entry gates (min)"],
            ["Boarding_Wait_Min", "Waiting time on the platform before boarding (min)"],
            ["Total_Time_Min", "Total time in system = sum of the five component times (min)"],
        ],
        widths=[1.7, 4.6],
    )

    doc.add_heading("3.2 Sample Rows", level=2)
    make_table(
        doc,
        ["ID", "Arrival", "Type", "C-Wait", "C-Serv", "G-Wait", "G-Serv", "B-Wait", "Total"],
        [
            ["P001", "07:00:02", "Online", "0.00", "0.00", "0.65", "1.47", "8.53", "10.65"],
            ["P002", "07:00:04", "Normal", "0.00", "0.88", "0.19", "0.76", "8.88", "10.71"],
            ["P003", "07:00:12", "Online", "0.00", "0.00", "0.33", "0.41", "1.00", "1.74"],
            ["P004", "07:00:25", "Online", "0.00", "0.00", "0.93", "0.71", "1.56", "3.19"],
            ["P005", "07:00:46", "Normal", "0.00", "0.36", "0.05", "0.85", "1.40", "2.67"],
        ],
        widths=[0.55, 0.9, 0.8, 0.7, 0.7, 0.7, 0.7, 0.7, 0.7],
    )

    doc.add_heading("3.3 Summary Statistics", level=2)
    doc.add_paragraph("All 400 passengers (times in minutes):")
    make_table(
        doc,
        ["Measure", "Counter Wait", "Counter Service", "Gate Wait", "Gate Service", "Boarding Wait", "Total Time"],
        [
            ["Mean", "0.55", "0.66", "0.49", "0.40", "5.05", "7.14"],
            ["Std dev", "1.70", "1.27", "0.51", "0.40", "4.72", "5.36"],
            ["Min", "0.00", "0.00", "0.05", "0.05", "1.00", "1.10"],
            ["Median", "0.00", "0.00", "0.33", "0.29", "3.52", "5.77"],
            ["Max", "8.49", "13.94", "3.14", "3.10", "29.81", "31.80"],
        ],
        widths=[0.9, 1.0, 1.1, 1.0, 1.0, 1.1, 1.0],
    )
    doc.add_paragraph("Normal passengers only (the class that uses the counters; n = 180):")
    make_table(
        doc,
        ["Measure", "Counter Wait", "Counter Service", "Total Time"],
        [
            ["Mean", "1.22", "1.46", "8.67"],
            ["Std dev", "2.36", "1.56", "5.61"],
            ["Median", "0.00", "1.08", "7.26"],
            ["Max", "8.49", "13.94", "31.80"],
        ],
        widths=[1.0, 1.5, 1.8, 1.2],
    )
    doc.add_paragraph(
        "The simulated mean counter service time of 1.46 min corresponds to a service "
        "rate of about μ = 41.2 passengers/hour (target 40), and the simulated mean "
        "counter wait of 1.22 min is close to the analytical M/M/c value of Wq = 1.14 "
        "min, which confirms that the data set is consistent with the model."
    )

    doc.add_heading("4. Repository and Next Steps", level=1)
    doc.add_paragraph(
        f"The complete project (data set, scripts, figures and this document) is stored "
        f"in a git repository: {GIT_URL}."
    )
    doc.add_paragraph(
        "In the next phase the analytical M/M/c (Erlang-C) model will be applied to "
        "these data to compute utilisation, waiting times and queue lengths, followed "
        "by a what-if capacity analysis and the final report."
    )

    out = DOCS / "Submission_1_Problem_and_Dataset.docx"
    doc.save(out)
    return out


def build_submission_2():
    doc = Document()
    add_title_block(doc, "Submission 2 – Final Report")

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "This report models passenger flow through the ticket concourse of a major "
        "railway station during the morning peak period (07:00 – 09:00) using M/M/c "
        "queueing theory (Erlang-C), Little's Law and a passenger-level discrete-event "
        "simulation. A simulated data set of 400 passengers (180 Normal, 220 Online) "
        "is used to estimate utilisation, waiting times and queue lengths. The analysis "
        "identifies the ticket counters as the bottleneck of the system: at the baseline "
        "peak load the utilisation is ρ = 0.75 and the expected wait Wq = 1.14 min, "
        "rising to ρ = 0.9167 and Wq = 5.08 min in a busy morning. The results show that "
        "adding two counters (c = 5) reduces the wait to 0.12 min, while faster service "
        "(μ = 50/hour) reduces it to 0.81 min. The simulated data set validates the "
        "analytical model (simulated mean wait 1.22 min vs analytical 1.14 min)."
    )

    doc.add_heading("1. Introduction and System Description", level=1)
    doc.add_paragraph(
        "Railway stations in the morning peak behave as coupled queueing networks. "
        "Passengers arrive stochastically, join queues at ticket counters and gates, and "
        "accumulate on platforms before boarding. The ticket counters serve only Normal "
        "ticket passengers while Online ticket passengers bypass them, so the counter "
        "subsystem carries a class-dependent and highly concentrated load."
    )
    doc.add_paragraph(
        "This project studies the flow of passengers through the stages: ticket "
        "acquisition (Normal class only), entry gates, platform and boarding. The "
        "performance objectives are to identify the bottleneck stage, quantify waiting "
        "times and queue lengths, establish the capacity limit of the current "
        "configuration, and evaluate what-if interventions that reduce congestion."
    )

    doc.add_heading("2. Modelling Approach and Assumptions", level=1)
    doc.add_paragraph(
        "The ticket-counter subsystem is modelled as an M/M/c queue: arrivals follow a "
        "Poisson process with rate λ, service times are exponential with rate μ per "
        "counter, and c identical counters operate in parallel with a FIFO discipline. "
        "The Erlang-C formula provides the probability of waiting and the expected "
        "waiting time, and Little's Law (Lq = λ × Wq) is used as a consistency check. "
        "The entry gates are modelled as an independent M/M/c queue for comparison."
    )
    doc.add_paragraph("Key assumptions:")
    for item in [
        "Arrivals are Poisson; the peak period is modelled as 2 hours of steady demand (07:00 – 09:00).",
        "Service times at the counters are exponential with mean 1/μ = 1.5 min (μ = 40/hour).",
        "The queue discipline is FIFO and there is no balking or reneging.",
        "The passenger-level data set is generated by a discrete-event simulation consistent with these assumptions and is simulated, not real observations.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2.1 Parameters", level=2)
    make_table(
        doc,
        ["Parameter", "Symbol", "Value"],
        [
            ["Normal arrival rate", "λ", "90 passengers/hour (180 / 2 h)"],
            ["Online arrival rate", "λ", "110 passengers/hour (220 / 2 h)"],
            ["Counter service rate", "μ", "40 passengers/hour per counter"],
            ["Number of counters", "c", "3 (baseline)"],
            ["Entry gate service rate", "μ", "100 passengers/hour per gate"],
            ["Number of gates", "c", "3"],
        ],
        widths=[2.2, 0.8, 3.3],
    )

    doc.add_heading("3. Data and Methodology", level=1)
    doc.add_paragraph(
        "The data set (data/passenger_flow_peak_hours.csv) contains 400 passengers: "
        "180 Normal (45.0%) and 220 Online (55.0%). It was produced by a FIFO M/M/c "
        "discrete-event simulation (c = 3, μ = 40/hour) with Poisson arrivals sampled as "
        "order statistics of a uniform distribution, giving an exact 07:00 – 09:00 "
        "window and a reproducible seed. Each passenger record captures counter wait, "
        "counter service, gate wait, gate service, boarding wait and total time."
    )
    doc.add_paragraph(
        "The scenario analysis is run as a controlled experiment: 20 independent "
        "2-hour replications of each scenario, same random seed per replication, with "
        "one parameter changed at a time. The analytical M/M/c values and the simulated "
        "means are compared to validate the model."
    )

    doc.add_heading("4. Analysis and Results", level=1)
    doc.add_heading("4.1 Baseline (Normal Peak)", level=2)
    make_table(
        doc,
        ["Measure", "Formula", "Value"],
        [
            ["Utilisation", "ρ = λ/(c·μ) = 90/(3×40)", "0.75"],
            ["Probability of waiting", "Erlang-C P(wait)", "0.626"],
            ["Expected queue length", "Lq", "1.703 passengers"],
            ["Expected wait", "Wq = Lq/λ", "1.14 min"],
        ],
        widths=[2.2, 2.5, 1.6],
    )
    doc.add_paragraph(
        "Little's Law check: Lq = λ × Wq = 90 × 1.14/60 = 1.71 ≈ 1.703 passengers "
        "(matches within rounding)."
    )

    doc.add_heading("4.2 Busy Morning (High Load)", level=2)
    make_table(
        doc,
        ["Measure", "Value"],
        [
            ["Arrival rate λ", "110 passengers/hour"],
            ["Utilisation ρ", "0.9167"],
            ["Expected wait Wq", "5.08 min"],
            ["Status", "Stable, but heavily congested (ρ close to 1)"],
        ],
        widths=[2.2, 4.0],
    )
    doc.add_paragraph(
        "At ρ = 0.9167 the system is only just stable; with c = 3 counters the capacity "
        "limit is λ < 120 passengers/hour."
    )

    doc.add_heading("4.3 What-if Interventions", level=2)
    make_table(
        doc,
        ["Scenario", "λ", "μ", "c", "ρ", "Wq (min)"],
        [
            ["Baseline (Normal Peak)", "90", "40", "3", "0.750", "1.14"],
            ["Busy Morning", "110", "40", "3", "0.9167", "5.08"],
            ["Busy + 5 Counters", "110", "40", "5", "0.55", "0.12"],
            ["Busy + Faster Service", "110", "50", "3", "0.7333", "0.81"],
        ],
        widths=[2.1, 0.7, 0.7, 0.7, 0.9, 1.0],
    )
    doc.add_paragraph(
        "Adding two counters (c = 5) reduces the busy-morning wait from 5.08 min to "
        "0.12 min, and increasing the service rate to μ = 50/hour reduces it to "
        "0.81 min. Both interventions bring utilisation below 0.75 and restore stable, "
        "short queues. The entry gates (ρ = 0.667, Wq = 0.27 min) are not the "
        "bottleneck."
    )

    doc.add_heading("4.4 Model Validation", level=2)
    doc.add_paragraph(
        "The simulated data set reproduces the analytical model: the simulated mean "
        "counter service time is 1.46 min (μ ≈ 41.2/hour vs target 40), and the "
        "simulated mean counter wait of 1.22 min is within 7.2% of the analytical "
        "Wq = 1.14 min. The finite 2-hour window starts empty, so simulated waits "
        "include the transient warm-up and sit at or below the steady-state M/M/c "
        "values; the gap grows with utilisation (e.g. Busy Morning: simulated ≈ 2.99 "
        "min vs steady-state 5.08 min)."
    )

    doc.add_heading("4.5 Sensitivity Analysis", level=2)
    doc.add_paragraph(
        "Sweeping the arrival rate from λ = 70 to 130 passengers/hour shows that Wq "
        "grows sharply as ρ approaches 1. With three counters the system becomes "
        "unstable above λ = 120; with five counters the same load stays comfortably "
        "within stable utilisation."
    )

    doc.add_heading("5. Visualisations", level=1)
    doc.add_paragraph("The following figures are generated by the analysis pipeline (see outputs/):")
    for num, cap in [
        ("01", "Utilisation of ticket counters (baseline and busy) and entry gates, with the high-congestion threshold."),
        ("02", "Effect of the number of counters on the expected wait Wq under λ = 110/hour."),
        ("03", "Arrival intensity across 07:00 – 09:00 (passengers per 10-minute bin)."),
        ("04", "Counter wait, gate wait and total time for Normal vs Online passengers."),
        ("05", "Sensitivity of Wq to the arrival rate for c = 3, 4 and 5 counters."),
        ("06", "Average counter waiting time per 10-minute bin, with the analytical Wq as reference."),
        ("07", "System flow diagram: Normal passengers queue at counters, Online passengers bypass them; both proceed through gates, platform and boarding."),
    ]:
        doc.add_paragraph(f"Figure {num}: {cap}", style="List Number")
    for f in sorted(OUTPUTS.glob("*.png")):
        try:
            doc.add_picture(str(f), width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = caption.add_run(f.stem)
            run.italic = True
            run.font.size = Pt(9)
        except Exception as e:
            doc.add_paragraph(f"(Figure {f.name} could not be embedded: {e})")

    doc.add_heading("6. Limitations and Future Work", level=1)
    for item in [
        "The data are simulated, not real station observations; the results should be confirmed with measured arrival and service data.",
        "The 2-hour window introduces a transient effect that the steady-state M/M/c model does not capture; a time-varying (transient) queueing model would remove this gap.",
        "Assumptions of Poisson arrivals, exponential service and no balking simplify real passenger behaviour.",
        "Only two passenger classes are distinguished; future work could add season tickets, family groups and platform capacity constraints.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("7. Conclusion", level=1)
    doc.add_paragraph(
        "The ticket counters are the bottleneck of the peak-period passenger flow. At "
        "the baseline load of λ = 90/hour the expected wait is 1.14 min at ρ = 0.75, "
        "but a busy morning at λ = 110/hour pushes utilisation to 0.9167 and the wait "
        "to 5.08 min. Adding two counters is the most effective intervention (Wq = "
        "0.12 min), and the analytical model is validated by the simulated data set. "
        "The project demonstrates a complete performance-modelling chain from data, "
        "through an analytical M/M/c model and validation, to capacity recommendations."
    )

    doc.add_heading("8. References", level=1)
    for ref in [
        "[1] D. G. Kendall, \"Stochastic processes occurring in the theory of queues and their analysis by the method of the imbedded Markov chain,\" The Annals of Mathematical Statistics, vol. 24, no. 3, pp. 338–354, 1953.",
        "[2] A. K. Erlang, \"Solution of some problems in the theory of probabilities of significance in automatic telephone exchanges,\" The Post Office Electrical Engineers' Journal, vol. 10, pp. 189–197, 1917.",
        "[3] J. D. C. Little, \"A proof for the queuing formula: L = λW,\" Operations Research, vol. 9, no. 3, pp. 383–387, 1961.",
        "[4] D. Gross, J. F. Shortle, J. M. Thompson, and C. M. Harris, Fundamentals of Queueing Theory, 4th ed. Hoboken, NJ, USA: Wiley, 2008.",
        "[5] L. Kleinrock, Queueing Systems, Volume 1: Theory. New York, NY, USA: Wiley, 1975.",
        "[6] K. S. Trivedi, Probability and Statistics with Reliability, Queuing, and Computer Science Applications, 2nd ed. New York, NY, USA: Wiley, 2002.",
    ]:
        doc.add_paragraph(ref, style="List Bullet")

    doc.add_heading("Repository", level=1)
    doc.add_paragraph(
        f"The full project (data set, scripts, results and this report) is available "
        f"in the git repository: {GIT_URL}."
    )

    out = DOCS / "Submission_2_Final_Report.docx"
    doc.save(out)
    return out


def main():
    DOCS.mkdir(exist_ok=True)
    s1 = build_submission_1()
    print(f"Created {s1}")
    s2 = build_submission_2()
    print(f"Created {s2}")


if __name__ == "__main__":
    main()
